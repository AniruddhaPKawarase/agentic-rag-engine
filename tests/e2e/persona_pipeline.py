"""Persona-based end-to-end test pipeline for Unified RAG Agent (Sandbox).

Reads questions from persona_questions.json (parsed from the persona docx) and
runs them concurrently against the sandbox ai6 gateway. Captures the same
columns as the 2026-04-13 baseline CSV plus persona/category context.

Artifacts per run, under test_results/persona_run_<timestamp>/:
- persona_run_results.csv     CSV matching baseline schema + persona/category
- persona_run_summary.json    Aggregate stats + full per-row records
- persona_run_summary.docx    Human-readable side-by-side report

Usage examples (from repo root):

  # Quick 10-question smoke test on project 7166
  python PROD_SETUP/unified-rag-agent/tests/e2e/persona_pipeline.py \
      --project-id 7166 --limit 10

  # Full 423-question run, 10 concurrent workers
  python PROD_SETUP/unified-rag-agent/tests/e2e/persona_pipeline.py \
      --project-id 7166 --concurrency 10

  # One persona only
  python PROD_SETUP/unified-rag-agent/tests/e2e/persona_pipeline.py \
      --project-id 7166 --persona "Project Manager"

  # Filter by category
  python PROD_SETUP/unified-rag-agent/tests/e2e/persona_pipeline.py \
      --project-id 7166 --category "Budget & Costs"
"""

from __future__ import annotations

import argparse
import concurrent.futures
import csv
import json
import sys
import time
from datetime import datetime
from pathlib import Path
from typing import Any

import requests


DEFAULT_BASE_URL = "https://ai6.ifieldsmart.com/rag"
DEFAULT_PROJECT_ID = 7166
DEFAULT_CONCURRENCY = 10
DEFAULT_TIMEOUT_S = 180

SCRIPT_DIR = Path(__file__).resolve().parent
# tests/e2e/ -> tests/ -> unified-rag-agent/
UNIFIED_RAG_ROOT = SCRIPT_DIR.parents[1]
DEFAULT_QUESTIONS = SCRIPT_DIR / "persona_questions.json"
RESULTS_ROOT = UNIFIED_RAG_ROOT / "test_results"


CSV_COLUMNS = [
    "persona", "category", "question", "answer", "confidence", "confidence_score",
    "agentic_confidence", "engine_used", "fallback_used", "model_used",
    "retrieval_count", "s3_path_count", "source_s3_paths_sample",
    "source_pdf_names_sample", "latency_seconds", "server_latency_ms",
    "total_tokens", "prompt_tokens", "completion_tokens", "http_status",
    "success", "error",
]


def load_questions(path: Path) -> list[dict[str, str]]:
    """Load the persona questions JSON written by the extractor."""
    if not path.exists():
        raise FileNotFoundError(
            f"Question file not found: {path}. Re-run the extractor from the docx."
        )
    with path.open("r", encoding="utf-8") as fh:
        data = json.load(fh)
    return data


def filter_questions(
    questions: list[dict[str, str]],
    persona: str | None,
    category: str | None,
    limit: int | None,
) -> list[dict[str, str]]:
    out = list(questions)
    if persona:
        out = [q for q in out if q["persona"].lower() == persona.lower()]
    if category:
        out = [q for q in out if q["category"].lower() == category.lower()]
    if limit and limit > 0:
        out = out[:limit]
    return out


def run_single(
    session: requests.Session,
    base_url: str,
    project_id: int,
    q_record: dict[str, str],
    timeout: int,
) -> dict[str, Any]:
    """Call /query for a single question and normalise the response."""
    start = time.monotonic()
    status = 0
    body: dict[str, Any] = {}
    err: str | None = None
    try:
        r = session.post(
            f"{base_url.rstrip('/')}/query",
            json={"query": q_record["question"], "project_id": project_id},
            timeout=timeout,
        )
        status = r.status_code
        body = r.json() if r.content else {}
    except requests.Timeout:
        err = f"timeout after {timeout}s"
    except requests.RequestException as exc:
        err = f"request failed: {exc}"
    except ValueError as exc:
        err = f"non-json response: {exc}"

    elapsed = round(time.monotonic() - start, 2)
    token_usage = body.get("token_usage") or {}
    sd = body.get("source_documents") or []
    s3_paths = [s.get("s3_path", "") for s in sd if s.get("s3_path")]
    pdf_names = [s.get("pdf_name", "") for s in sd if s.get("pdf_name")]

    return {
        "persona": q_record["persona"],
        "category": q_record["category"],
        "question": q_record["question"],
        "answer": body.get("answer", "") or "",
        "confidence": body.get("confidence"),
        "confidence_score": body.get("confidence_score"),
        "agentic_confidence": body.get("agentic_confidence"),
        "engine_used": body.get("engine_used"),
        "fallback_used": body.get("fallback_used"),
        "model_used": body.get("model_used"),
        "retrieval_count": body.get("retrieval_count") or len(sd),
        "s3_path_count": len(s3_paths),
        "source_s3_paths_sample": "; ".join(s3_paths[:3]),
        "source_pdf_names_sample": "; ".join(pdf_names[:3]),
        "latency_seconds": elapsed,
        "server_latency_ms": body.get("processing_time_ms"),
        "total_tokens": token_usage.get("total_tokens"),
        "prompt_tokens": token_usage.get("prompt_tokens"),
        "completion_tokens": token_usage.get("completion_tokens"),
        "http_status": status,
        "success": (status == 200 and not err and bool(body.get("answer"))),
        "error": err or body.get("error"),
        "_source_documents": sd,
    }


def run_parallel(
    base_url: str,
    project_id: int,
    questions: list[dict[str, str]],
    concurrency: int,
    timeout: int,
) -> list[dict[str, Any]]:
    """Dispatch questions through a thread pool and return ordered results."""
    results: list[dict[str, Any] | None] = [None] * len(questions)
    session = requests.Session()
    with concurrent.futures.ThreadPoolExecutor(max_workers=concurrency) as pool:
        futures = {
            pool.submit(run_single, session, base_url, project_id, q, timeout): i
            for i, q in enumerate(questions)
        }
        done = 0
        total = len(futures)
        for fut in concurrent.futures.as_completed(futures):
            i = futures[fut]
            results[i] = fut.result()
            done += 1
            q_short = questions[i]["question"][:70].replace("\n", " ")
            r = results[i] or {}
            print(
                f"[{done:>4d}/{total}] {r.get('engine_used') or '?':<11} "
                f"fb={str(r.get('fallback_used'))[:5]:<5} "
                f"s3={r.get('s3_path_count'):>2} "
                f"t={r.get('latency_seconds')}s  {q_short}",
                flush=True,
            )
    return [r for r in results if r is not None]


def write_csv(results: list[dict[str, Any]], path: Path) -> None:
    path.parent.mkdir(parents=True, exist_ok=True)
    with path.open("w", encoding="utf-8", newline="") as fh:
        writer = csv.DictWriter(fh, fieldnames=CSV_COLUMNS, extrasaction="ignore")
        writer.writeheader()
        for r in results:
            writer.writerow({k: r.get(k, "") for k in CSV_COLUMNS})


def write_json(results: list[dict[str, Any]], summary: dict[str, Any], path: Path) -> None:
    path.parent.mkdir(parents=True, exist_ok=True)
    # Strip large inner fields we don't need in the JSON summary, but keep source docs
    payload = {"summary": summary, "results": results}
    path.write_text(json.dumps(payload, indent=2, default=str, ensure_ascii=False), encoding="utf-8")


def write_docx(results: list[dict[str, Any]], summary: dict[str, Any], path: Path) -> None:
    try:
        from docx import Document
    except ImportError:
        print("python-docx not installed; skipping DOCX report.")
        return
    doc = Document()
    doc.add_heading("Persona-Based RAG Pipeline Run", level=1)
    doc.add_paragraph(f"Timestamp: {summary['timestamp']}")
    doc.add_paragraph(f"Base URL: {summary['base_url']}")
    doc.add_paragraph(f"Project ID: {summary['project_id']}")
    doc.add_paragraph(f"Questions: {summary['num_questions']}")
    doc.add_paragraph(f"Successful: {summary['successful']}  |  Failed: {summary['failed']}")
    doc.add_paragraph(f"Avg latency: {summary['avg_latency_s']}s  |  Total tokens: {summary['total_tokens']}")
    doc.add_paragraph(f"Fallback triggered: {summary['fallback_triggered_count']}/{summary['num_questions']}")
    doc.add_paragraph(f"Null s3_path queries: {summary['null_s3_path_queries']}/{summary['num_questions']}")
    doc.add_paragraph(f"Engine distribution: {summary['engine_distribution']}")
    doc.add_paragraph("")

    for i, r in enumerate(results, 1):
        doc.add_heading(f"Q{i}. [{r['persona']}/{r['category']}] {r['question'][:90]}", level=2)
        p = doc.add_paragraph()
        p.add_run("Engine: ").bold = True
        p.add_run(str(r.get("engine_used")))
        p.add_run("  |  Fallback: ").bold = True
        p.add_run(str(r.get("fallback_used")))
        p.add_run("  |  Conf: ").bold = True
        p.add_run(f"{r.get('confidence')}/{r.get('confidence_score')}")
        p.add_run("  |  Sources: ").bold = True
        p.add_run(str(r.get("retrieval_count")))
        p.add_run("  |  s3_paths: ").bold = True
        p.add_run(str(r.get("s3_path_count")))

        doc.add_paragraph("Answer:", style="Intense Quote")
        doc.add_paragraph(r.get("answer") or "[empty]")

        if r.get("_source_documents"):
            doc.add_paragraph("Top 3 sources:", style="Intense Quote")
            for s in r["_source_documents"][:3]:
                doc.add_paragraph(
                    f"- {s.get('pdf_name') or s.get('file_name') or '?'}  ({s.get('s3_path') or 'NO S3_PATH'})"
                )
    doc.save(str(path))


def summarise(results: list[dict[str, Any]], base_url: str, project_id: int) -> dict[str, Any]:
    successful = sum(1 for r in results if r.get("success"))
    failed = len(results) - successful
    avg_latency = round(
        sum((r.get("latency_seconds") or 0) for r in results) / max(len(results), 1), 2
    )
    total_tokens = sum((r.get("total_tokens") or 0) for r in results)
    fb = sum(1 for r in results if r.get("fallback_used"))
    null_s3 = sum(1 for r in results if (r.get("s3_path_count") or 0) == 0)
    engines = {
        e: sum(1 for r in results if r.get("engine_used") == e)
        for e in {r.get("engine_used") for r in results if r.get("engine_used")}
    }
    return {
        "timestamp": datetime.now().isoformat(),
        "base_url": base_url,
        "project_id": project_id,
        "num_questions": len(results),
        "successful": successful,
        "failed": failed,
        "avg_latency_s": avg_latency,
        "total_tokens": total_tokens,
        "fallback_triggered_count": fb,
        "null_s3_path_queries": null_s3,
        "engine_distribution": engines,
    }


def main() -> int:
    parser = argparse.ArgumentParser(description="Persona-based RAG pipeline test runner")
    parser.add_argument("--base-url", default=DEFAULT_BASE_URL, help="RAG base URL (default: %(default)s)")
    parser.add_argument("--project-id", type=int, default=DEFAULT_PROJECT_ID, help="Project ID (default: %(default)s)")
    parser.add_argument("--questions", type=Path, default=DEFAULT_QUESTIONS, help="Path to persona_questions.json")
    parser.add_argument("--persona", default=None, help="Filter: e.g. 'Project Manager'")
    parser.add_argument("--category", default=None, help="Filter: e.g. 'Budget & Costs'")
    parser.add_argument("--limit", type=int, default=None, help="Cap total questions")
    parser.add_argument("--concurrency", type=int, default=DEFAULT_CONCURRENCY, help="Parallel workers (default: %(default)s)")
    parser.add_argument("--timeout", type=int, default=DEFAULT_TIMEOUT_S, help="Per-request timeout seconds (default: %(default)s)")
    parser.add_argument("--out-dir", type=Path, default=None, help="Override output folder")
    args = parser.parse_args()

    all_q = load_questions(args.questions)
    filtered = filter_questions(all_q, args.persona, args.category, args.limit)
    if not filtered:
        print("No questions matched filters; exiting.", file=sys.stderr)
        return 2

    ts = datetime.now().strftime("%Y%m%d_%H%M%S")
    out_dir = args.out_dir or (RESULTS_ROOT / f"persona_run_{ts}")
    out_dir.mkdir(parents=True, exist_ok=True)

    print(
        f"Running {len(filtered)} questions on project {args.project_id} "
        f"via {args.base_url} with concurrency={args.concurrency}.\n"
        f"Output: {out_dir}",
        flush=True,
    )

    results = run_parallel(
        base_url=args.base_url,
        project_id=args.project_id,
        questions=filtered,
        concurrency=args.concurrency,
        timeout=args.timeout,
    )

    summary = summarise(results, args.base_url, args.project_id)
    write_csv(results, out_dir / "persona_run_results.csv")
    write_json(results, summary, out_dir / "persona_run_summary.json")
    write_docx(results, summary, out_dir / "persona_run_summary.docx")

    print("\n=== Summary ===")
    for k, v in summary.items():
        print(f"  {k}: {v}")
    print(f"\nArtifacts saved to: {out_dir}")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
