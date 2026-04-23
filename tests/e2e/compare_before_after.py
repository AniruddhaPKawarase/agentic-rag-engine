"""Produce a side-by-side before/after report from two runs.

Inputs: directories from reproduce_manager_failures.py (sandbox) and
local_fix_harness.py (local). Writes a JSON + DOCX side-by-side report.

Usage:
    python PROD_SETUP/unified-rag-agent/tests/e2e/compare_before_after.py \
        --before test_results/manager_before_fixes_GETprobe_20260422_151116 \
        --after test_results/local_after_fixes_local_20260422_152341 \
        --out test_results/comparison_fixes_8_5_<ts>
"""

from __future__ import annotations

import argparse
import json
from datetime import datetime
from pathlib import Path
from typing import Any


def load(run_dir: Path) -> tuple[list[dict[str, Any]], dict[str, Any]]:
    records = json.loads((run_dir / "records.json").read_text(encoding="utf-8"))
    summary = json.loads((run_dir / "summary.json").read_text(encoding="utf-8"))
    return records, summary


def key(r: dict[str, Any]) -> tuple[int, str]:
    return (int(r.get("project_id") or 0), str(r.get("question_id")))


def main() -> int:
    p = argparse.ArgumentParser()
    p.add_argument("--before", required=True, type=Path)
    p.add_argument("--after", required=True, type=Path)
    p.add_argument("--out", required=True, type=Path)
    args = p.parse_args()

    bef_rec, bef_sum = load(args.before)
    aft_rec, aft_sum = load(args.after)

    aft_by_key = {key(r): r for r in aft_rec}

    rows: list[dict[str, Any]] = []
    for bef in bef_rec:
        k = key(bef)
        aft = aft_by_key.get(k, {})
        bef_sumry = bef.get("summary") or bef  # reproduce script uses 'summary' wrap
        rows.append({
            "project_id": bef.get("project_id"),
            "question_id": bef.get("question_id"),
            "question": bef.get("question"),
            "before_engine": bef_sumry.get("engine_used"),
            "after_engine": aft.get("engine_used"),
            "before_fallback": bef_sumry.get("fallback_used"),
            "after_fallback": aft.get("fallback_used"),
            "before_confidence": bef_sumry.get("confidence"),
            "after_confidence": aft.get("confidence"),
            "before_sources": bef_sumry.get("source_documents_total"),
            "after_sources": aft.get("source_documents_total"),
            "before_signed_urls": bef_sumry.get("source_documents_with_s3_path"),
            "after_signed_urls": aft.get("source_documents_with_signed_url"),
            "before_url_ok_ratio": bef.get("url_probe_ok_ratio"),
            "after_url_ok_ratio": aft.get("url_probe_ok_ratio"),
            "before_answer_len": bef_sumry.get("answer_length"),
            "after_answer_len": aft.get("answer_length"),
            "before_answer": (bef.get("answer") or "")[:500],
            "after_answer": (aft.get("answer") or "")[:500],
        })

    delta = {
        "timestamp": datetime.now().isoformat(),
        "before_label": bef_sum.get("label"),
        "after_label": aft_sum.get("label"),
        "before_base_url": bef_sum.get("base_url"),
        "after_base_url": "local in-process",
        "before_queries_with_broken_url": bef_sum.get("queries_with_any_broken_url"),
        "after_queries_with_broken_url": aft_sum.get("queries_with_any_broken_url_GET"),
        "before_queries_with_zero_sources": bef_sum.get("queries_with_zero_sources"),
        "after_queries_with_zero_sources": aft_sum.get("queries_with_zero_sources"),
        "before_engine_distribution": bef_sum.get("engine_distribution"),
        "after_engine_distribution": aft_sum.get("engine_distribution"),
        "before_avg_latency_s": bef_sum.get("avg_elapsed_s"),
        "after_avg_latency_s": aft_sum.get("avg_elapsed_s"),
        "improvements_claimed": [
            "Fix #8: signed download URLs for all sources (traditional + agentic)",
            "Fix #5: parent-document retrieval for specs reduces 'full text not available' dodges",
        ],
    }

    args.out.mkdir(parents=True, exist_ok=True)
    (args.out / "comparison.json").write_text(
        json.dumps({"aggregate": delta, "rows": rows}, indent=2, ensure_ascii=False),
        encoding="utf-8",
    )

    # DOCX
    try:
        from docx import Document
        doc = Document()
        doc.add_heading("Manager Failures — Before vs After Fixes", level=1)
        doc.add_paragraph(f"Generated: {delta['timestamp']}")
        doc.add_paragraph(f"Before: {delta['before_label']} ({delta['before_base_url']})")
        doc.add_paragraph(f"After:  {delta['after_label']} (local in-process, no VM deploy)")
        doc.add_heading("Aggregate deltas", level=2)
        for k in [
            "before_queries_with_broken_url", "after_queries_with_broken_url",
            "before_queries_with_zero_sources", "after_queries_with_zero_sources",
            "before_engine_distribution", "after_engine_distribution",
            "before_avg_latency_s", "after_avg_latency_s",
        ]:
            doc.add_paragraph(f"{k}: {delta[k]}")

        for r in rows:
            doc.add_heading(f"[{r['project_id']}] {r['question_id']}  {r['question'][:80]}", level=2)
            p = doc.add_paragraph()
            p.add_run("Before: ").bold = True
            p.add_run(
                f"engine={r['before_engine']} fb={r['before_fallback']} "
                f"conf={r['before_confidence']} srcs={r['before_sources']} "
                f"signed_urls={r['before_signed_urls']} url_ok={r['before_url_ok_ratio']} "
                f"len={r['before_answer_len']}"
            )
            p = doc.add_paragraph()
            p.add_run("After:  ").bold = True
            p.add_run(
                f"engine={r['after_engine']} fb={r['after_fallback']} "
                f"conf={r['after_confidence']} srcs={r['after_sources']} "
                f"signed_urls={r['after_signed_urls']} url_ok={r['after_url_ok_ratio']} "
                f"len={r['after_answer_len']}"
            )
            doc.add_paragraph("Before answer:", style="Intense Quote")
            doc.add_paragraph(r["before_answer"] or "[empty]")
            doc.add_paragraph("After answer:", style="Intense Quote")
            doc.add_paragraph(r["after_answer"] or "[empty]")
        doc.save(str(args.out / "comparison.docx"))
    except ImportError:
        print("python-docx missing, DOCX skipped")

    print(f"Comparison written to {args.out}")
    print(json.dumps(delta, indent=2))
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
